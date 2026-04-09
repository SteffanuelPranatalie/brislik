import streamlit as st
import pandas as pd
import json
import io
from datetime import datetime

# Import library pendukung
try:
    from docx import Document
    from docx.enum.section import WD_ORIENTATION
    from docx.shared import Inches, Pt
    from fpdf import FPDF
except ImportError:
    st.error("Pustaka pendukung Word/PDF belum terinstal. Jalankan: python -m pip install python-docx fpdf2 openpyxl")

# 1. Konfigurasi Halaman
st.set_page_config(
    page_title="BRISLIK",
    page_icon="https://upload.wikimedia.org/wikipedia/commons/6/6d/BRI_2025.png?20251217000202",
    layout="wide"
)

# 2. Desain Dashboard
st.markdown("""
    <style>
    .stApp { background-color: var(--background-color); }
    .main-title { color: #003366 !important; font-size: 30px; font-weight: 800; border-bottom: 4px solid #3399FF; padding-bottom: 10px; margin-bottom: 25px; }
    .box-container { padding: 20px; border-radius: 12px; margin-bottom: 20px; height: 380px !important; overflow-y: auto; }
    .identitas-bg { background-color: #F8F9FA !important; border: 2px solid #D1D5DB !important; }
    .audit-bg { background-color: #FFF9C4 !important; border: 2px solid #FBC02D !important; }
    .inner-header { color: #003366 !important; font-size: 18px; font-weight: 800; margin-bottom: 15px; border-bottom: 1px solid rgba(0,0,0,0.1); text-transform: uppercase; }
    .lbl { color: #6B7280 !important; font-size: 11px; font-weight: 800; text-transform: uppercase; margin-bottom: 1px; }
    .val { color: #111827 !important; font-size: 14px; font-weight: 700; margin-bottom: 8px; line-height: 1.3; }
    .table-header { color: #003366 !important; font-size: 20px; font-weight: 700; margin-top: 20px; margin-bottom: 15px; border-left: 6px solid #3399FF; padding-left: 15px; }
    .blue-header thead tr th { background-color: #0000FF !important; color: white !important; }
    @media (prefers-color-scheme: dark) {
        .main-title, .table-header { color: #99CCFF !important; }
        .inner-header { color: #003366 !important; }
        .val { color: #111827 !important; }
        .lbl { color: #4B5563 !important; }
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown('<div class="main-title">📊 BRISLIK Rekapitulasi & Audit</div>', unsafe_allow_html=True)

# --- FUNGSI HELPER ---
def to_float(val):
    if not val or val == '-': return 0.0
    clean_val = str(val).upper().replace('RP', '').replace('.', '').replace(' ', '').strip()
    clean_val = clean_val.split(',')[0]
    try: return float(clean_val)
    except ValueError: return 0.0

def format_rupiah(val):
    if isinstance(val, str) and "Rp" in val: return val
    try: return "Rp " + f"{int(to_float(val)):,}".replace(",", ".")
    except: return "Rp 0"

def format_date(date_str):
    if not date_str or date_str in ["-", "null", ""]: return "-"
    try:
        dt = datetime.strptime(str(date_str)[:8], '%Y%m%d')
        return dt.strftime('%d-%m-%Y')
    except: return date_str

def safe_text(text):
    if not text: return "-"
    return str(text).replace("✔", "V").encode('ascii', 'ignore').decode('ascii')

# --- FUNGSI EKSPOR ---
def export_excel(id_info, aud_info, df):
    output = io.BytesIO()
    
    # Konversi menjadi integer mentah agar bisa di SUM oleh Excel
    plafon_num = int(to_float(aud_info['plafon']))
    baki_num = int(to_float(aud_info['baki']))
    
    summary_data = [
        ["IDENTITAS DEBITUR", ""],
        ["Nama Lengkap", id_info['nama']], ["NIK", id_info['nik']],
        ["Tempat/Tgl Lahir", f"{id_info['tmpt_lahir']}, {id_info['tgl_lahir']}"],
        ["Jenis Kelamin", id_info['jk']], ["NPWP", id_info['npwp']],
        ["Pekerjaan", id_info['pekerjaan']], ["Alamat", id_info['alamat']],
        ["", ""],
        ["SUMMARY AUDIT", ""],
        ["Skor Terburuk", f"Kolektabilitas {aud_info['skor']}"],
        ["Total Plafon", plafon_num], ["Total Kewajiban", baki_num],
        ["Utilisasi", aud_info['util']], ["Kreditur", f"{aud_info['total_kred']} Lembaga"],
        ["Posisi Data", aud_info['posisi']], ["Tanggal Laporan", id_info['tgl']],
        ["", ""]
    ]
    df_sum = pd.DataFrame(summary_data)
    
    df_export = df.copy()
    kolom_uang = ["PLAFON", "BAKI DEBET", "OS (Rp)", "Plafon Awal", "OS"]
    
    # Mengubah value ke number agar fungsi SUM di Excel aktif
    for col in kolom_uang:
        if col in df_export.columns:
            df_export[col] = df_export[col].apply(lambda x: int(to_float(x)))
            
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_sum.to_excel(writer, index=False, header=False, sheet_name='Audit')
        if not df_export.empty: 
            df_export.to_excel(writer, index=False, startrow=len(summary_data), sheet_name='Audit')
        
        # --- LOGIKA PENAMBAHAN FORMAT RUPIAH PADA EXCEL NUMBER ---
        workbook = writer.book
        worksheet = writer.sheets['Audit']
        
        # Format string dari OpenPyXL yang memberitahu Excel untuk menampilkan "Rp " di depan number
        rp_format = '"Rp "#,##0'
        
        # Menerapkan format di Summary (Baris ke-12 dan 13)
        worksheet.cell(row=12, column=2).number_format = rp_format
        worksheet.cell(row=13, column=2).number_format = rp_format
        
        # Menerapkan format ke kolom tabel rincian
        if not df_export.empty:
            for col_name in kolom_uang:
                if col_name in df_export.columns:
                    col_idx = df_export.columns.get_loc(col_name) + 1
                    for row_idx in range(len(summary_data) + 2, len(summary_data) + 2 + len(df_export)):
                        worksheet.cell(row=row_idx, column=col_idx).number_format = rp_format

    return output.getvalue()

def export_word(id_info, aud_info, df):
    doc = Document(); section = doc.sections[-1]; section.orientation = WD_ORIENTATION.LANDSCAPE
    doc.add_heading('LAPORAN REKAPITULASI & AUDIT BRISLIK', 0)
    doc.add_heading('IDENTITAS DEBITUR', level=1)
    doc.add_paragraph(f"Nama: {id_info['nama']}\nNIK: {id_info['nik']}\nTTL: {id_info['tmpt_lahir']}, {id_info['tgl_lahir']}\nJK: {id_info['jk']} | NPWP: {id_info['npwp']}\nPekerjaan: {id_info['pekerjaan']}\nAlamat: {id_info['alamat']}")
    doc.add_heading('SUMMARY AUDIT', level=1)
    doc.add_paragraph(f"Skor: Kolektabilitas {aud_info['skor']}\nTotal Plafon: {aud_info['plafon']}\nTotal Kewajiban: {aud_info['baki']}\nUtilisasi: {aud_info['util']} | Kreditur: {aud_info['total_kred']} Lembaga\nPosisi Data: {aud_info['posisi']}")
    if not df.empty:
        table = doc.add_table(rows=1, cols=len(df.columns)); table.style = 'Table Grid'
        for i, col in enumerate(df.columns): table.cell(0, i).text = col
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row): row_cells[i].text = safe_text(val)[:35]
    out = io.BytesIO(); doc.save(out); return out.getvalue()

def export_pdf(id_info, aud_info, df):
    pdf = FPDF('L', 'mm', 'A4'); pdf.add_page(); pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, f"LAPORAN AUDIT: {id_info['nama']}", ln=True, align='C'); pdf.ln(4)
    pdf.set_font("Helvetica", 'B', 9); pdf.cell(0, 6, "IDENTITAS DEBITUR", ln=True)
    pdf.set_font("Helvetica", size=8)
    pdf.cell(0, 5, safe_text(f"Nama: {id_info['nama']} | NIK: {id_info['nik']} | TTL: {id_info['tmpt_lahir']}, {id_info['tgl_lahir']}"), ln=True)
    pdf.cell(0, 5, safe_text(f"Pekerjaan: {id_info['pekerjaan']} | NPWP: {id_info['npwp']} | JK: {id_info['jk']}"), ln=True)
    pdf.ln(3); pdf.set_font("Helvetica", 'B', 9); pdf.cell(0, 6, "SUMMARY AUDIT", ln=True)
    pdf.set_font("Helvetica", size=8)
    pdf.cell(0, 5, safe_text(f"Skor: Kol {aud_info['skor']} | Plafon: {aud_info['plafon']} | Kewajiban: {aud_info['baki']}"), ln=True)
    pdf.cell(0, 5, safe_text(f"Utilisasi: {aud_info['util']} | Kreditur: {aud_info['total_kred']} Lembaga | Posisi: {aud_info['posisi']}"), ln=True)
    pdf.ln(5)
    
    if not df.empty:
        pdf.set_font("Helvetica", 'B', 6) 
        
        if "Restrukturisasi Iya" in df.columns:
            # Slik 3 (Egie): memiliki 11 Kolom
            w = [7, 25, 45, 35, 25, 18, 18, 20, 16, 30, 30]
        elif "OS (Rp)" in df.columns:
            # Slik 2 (Aldista): memiliki 14 Kolom
            w = [7, 20, 36, 25, 22, 22, 18, 18, 12, 12, 16, 10, 15, 15] 
        else:
            # Slik 1 (Default): memiliki 12 Kolom
            w = [7, 40, 30, 22, 28, 28, 20, 20, 15, 15, 15, 37] 
        
        for i, c in enumerate(df.columns): 
            pdf.cell(w[i], 8, safe_text(c)[:int(w[i]*0.9)], 1, 0, 'C')
        pdf.ln()
        
        pdf.set_font("Helvetica", size=5)
        for _, r in df.iterrows():
            for i, col in enumerate(df.columns): 
                val = safe_text(r[col])
                max_chars = int(w[i] * 0.95)
                pdf.cell(w[i], 7, val[:max_chars], 1, 0, 'L' if i in [1,2,3] else 'C')
            pdf.ln()
    return bytes(pdf.output())

# 3. Sidebar & Logika Utama
with st.sidebar:
    st.header("⚙️ Menu Utama")
    uploaded_files = st.file_uploader("Unggah File .txt iDEB", type=["txt"], accept_multiple_files=True)
    st.divider(); st.caption("Developed by Steffanuel Pranatalie")

if uploaded_files:
    for uploaded_file in uploaded_files:
        try:
            with st.expander(f"📁 Dashboard Debitur: {uploaded_file.name}", expanded=True):
                raw_content = uploaded_file.read().decode("utf-8-sig", errors="ignore")
                data = json.loads(raw_content.strip())
                ind = data.get('individual', {}); data_pokok = ind.get('dataPokokDebitur', [{}])[0]
                ringkasan = ind.get('ringkasanFasilitas', {}); header_info = data.get('header', {})
                
                nama_v = str(data_pokok.get('namaDebitur') or "-").upper()
                nik_v = str(data_pokok.get('noIdentitas', '-'))
                alamat_v = str(data_pokok.get('alamat', '-'))
                tgl_lahir_v = format_date(data_pokok.get('tanggalLahir', '-'))
                tmpt_lahir_v = str(data_pokok.get('tempatLahir', '-')).upper()
                jk_v = str(data_pokok.get('jenisKelaminKet', '-')).upper()
                npwp_v = str(data_pokok.get('npwp', '-'))
                pekerjaan_v = str(data_pokok.get('pekerjaanKet', '-'))
                tgl_laporan_v = format_date(header_info.get('tanggalHasil'))
                
                skor_v = ringkasan.get('kualitasTerburuk', '-')
                plafon_v = float(ringkasan.get('plafonEfektifTotal', 0))
                baki_v = float(ringkasan.get('bakiDebetTotal', 0))
                total_kred = sum([int(ringkasan.get(k, 0) or 0) for k in ['krediturBankUmum', 'krediturBPR/S', 'krediturLp', 'krediturLainnya']])
                posisi_data_v = str(ind.get('posisiDataTerakhir', '-'))

                fas_root = ind.get('fasilitas', {})
                all_fas = []
                for k in fas_root:
                    if isinstance(fas_root[k], list): all_fas.extend(fas_root[k])

                rows = []
                for i, f in enumerate(all_fas, 1):
                    histori_kolek = []
                    if f.get('kualitas'): histori_kolek.append(str(f.get('kualitas')))
                    for j in range(1, 25):
                        kunci_kol = f"tahunBulan{j:02d}Kol"
                        nilai_kol = f.get(kunci_kol)
                        if nilai_kol and str(nilai_kol).strip():
                            histori_kolek.append(str(nilai_kol).strip())
                    kolek_terburuk = max(histori_kolek) if histori_kolek else (f.get('kualitas') or '-')
                    
                    raw_p = str(f.get('jenisPenggunaanKet', '')).lower()
                    mapped_p = "KMK" if "modal kerja" in raw_p else ("Investasi" if "investasi" in raw_p else "Konsumsi")
                    original_p = f.get('jenisKreditPembiayaanKet') or f.get('jenisKreditKet', '-')
                    
                    tgl_mulai = format_date(f.get('tanggalMulai'))
                    tgl_jt = format_date(f.get('tanggalJatuhTempo'))
                    
                    plafon_awal_raw = f.get('nilaiProyek')
                    if not plafon_awal_raw or float(plafon_awal_raw) == 0:
                         plafon_awal_raw = f.get('plafon', 0)

                    rows.append({
                        "NO": i, "NAMA JASA KEUANGAN": (f.get('ljkKet') or '-').upper(), 
                        "JENIS_ORIGINAL": original_p, "JENIS_MAPPED": mapped_p,
                        "PLAFON": format_rupiah(f.get('plafon', 0)), 
                        "PLAFON_AWAL": format_rupiah(plafon_awal_raw), 
                        "BAKI DEBET": format_rupiah(f.get('bakiDebet', 0)),
                        "RAW_BAKI": float(f.get('bakiDebet', 0)), 
                        "TGL_MULAI": tgl_mulai, "JATUH_TEMPO": tgl_jt,
                        "KOL_TERAKHIR": str(f.get('kualitas') or '-'),
                        "KOL_TERBURUK": kolek_terburuk,
                        "BUNGA": f"{f.get('sukuBungaImbalan', '-')} %", "KONDISI": f.get('kondisiKet', '-'),
                        "RESTRUK": "Y" if f.get('tanggalRestrukturisasiAkhir') else "N"
                    })

                b_val = to_float(baki_v)
                p_val = to_float(plafon_v)
                util_v = (b_val / p_val * 100) if p_val > 0 else 0

                col_id, col_aud = st.columns(2)
                with col_id:
                    st.markdown(f"""<div class="box-container identitas-bg"><div class="inner-header">👤 Identitas Debitur</div>
                        <p class="lbl">Nama Lengkap</p><p class="val">{nama_v}</p>
                        <p class="lbl">NIK / NPWP</p><p class="val">{nik_v} / {npwp_v}</p>
                        <p class="lbl">TTL / Jenis Kelamin</p><p class="val">{tmpt_lahir_v}, {tgl_lahir_v} | {jk_v}</p>
                        <p class="lbl">Pekerjaan</p><p class="val">{pekerjaan_v}</p>
                        <p class="lbl">Alamat Lengkap</p><p class="val">{alamat_v}</p></div>""", unsafe_allow_html=True)
                with col_aud:
                    st.markdown(f"""<div class="box-container audit-bg"><div class="inner-header">🔍 Summary Audit</div>
                        <p class="lbl">Skor Terburuk</p><p class="val" style="color:red !important;">Kolektabilitas {skor_v}</p>
                        <p class="lbl">Total Plafon</p><p class="val">{format_rupiah(plafon_v)}</p>
                        <p class="lbl">Total Kewajiban</p><p class="val">{format_rupiah(baki_v)}</p>
                        <p class="lbl">Utilisasi & Kreditur</p><p class="val">{util_v:.2f}% | {total_kred} Lembaga</p>
                        <p class="lbl">Posisi Data Terakhir</p><p class="val">{posisi_data_v}</p></div>""", unsafe_allow_html=True)

                df_final = pd.DataFrame()
                if rows:
                    df_full = pd.DataFrame(rows)
                    st.markdown('<div class="table-header">PENGATURAN OUTPUT TABEL</div>', unsafe_allow_html=True)
                    
                    sel_format = st.radio(f"Pilih Tampilan ({uploaded_file.name}):", options=["slik 1 (Default)", "slik 2 (Aldista)", "slik 3 (Egie)"], horizontal=True, key=f"fmt_{uploaded_file.name}")
                    
                    c_f1, c_f2, c_f3, c_f4 = st.columns(4)
                    with c_f1: sel_bank = st.multiselect("Filter Bank", options=sorted(df_full['NAMA JASA KEUANGAN'].unique()), key=f"bank_{uploaded_file.name}")
                    with c_f2: sel_jenis_penggunaan = st.multiselect("Filter Jenis Penggunaan", options=sorted(df_full['JENIS_MAPPED'].unique()), key=f"jp_{uploaded_file.name}")
                    with c_f3: sel_jenis = st.multiselect("Filter Jenis", options=sorted(df_full['JENIS_ORIGINAL'].unique()), key=f"j_{uploaded_file.name}")
                    with c_f4: sel_kondisi = st.multiselect("Filter Kondisi", options=sorted(df_full['KONDISI'].unique()), key=f"kond_{uploaded_file.name}")

                    df_f = df_full.copy()
                    if sel_bank: df_f = df_f[df_f['NAMA JASA KEUANGAN'].isin(sel_bank)]
                    if sel_jenis_penggunaan: df_f = df_f[df_f['JENIS_MAPPED'].isin(sel_jenis_penggunaan)]
                    if sel_jenis: df_f = df_f[df_f['JENIS_ORIGINAL'].isin(sel_jenis)]
                    if sel_kondisi: df_f = df_f[df_f['KONDISI'].isin(sel_kondisi)]
                    df_f['NO'] = range(1, len(df_f) + 1)

                    st.markdown('<div class="table-header">RINCIAN FASILITAS DEBITUR</div>', unsafe_allow_html=True)
                    
                    if sel_format == "slik 3 (Egie)":
                        df_c = df_f.rename(columns={
                            "JENIS_MAPPED": "Jenis Penggunaan",
                            "NAMA JASA KEUANGAN": "Bank",
                            "BAKI DEBET": "OS",
                            "KOL_TERAKHIR": "Kol Terakhir",
                            "KOL_TERBURUK": "Kol Terburuk",
                            "BUNGA": "Suku Bunga"
                        })
                        df_c["Jumlah Hari Kol"] = "-" 
                        df_c["Restrukturisasi Iya"] = df_c["RESTRUK"].apply(lambda x: "✔" if x == "Y" else "")
                        df_c["Restrukturisasi Tidak"] = df_c["RESTRUK"].apply(lambda x: "✔" if x == "N" else "")
                        
                        cols_slik3 = ["NO", "Jenis Penggunaan", "Bank", "OS", "Kol Terakhir", "Kol Terburuk", "Jumlah Hari Kol", "Suku Bunga", "Restrukturisasi Iya", "Restrukturisasi Tidak"]
                        st.markdown('<div class="blue-header">', unsafe_allow_html=True); st.dataframe(df_c[cols_slik3], use_container_width=True, hide_index=True); st.markdown('</div>', unsafe_allow_html=True)
                        st.markdown(f"""<div style="background-color:#0000FF; color:white; padding:10px; font-weight:bold; text-align:center;">Total Outstanding: {format_rupiah(df_f['RAW_BAKI'].sum())}</div>""", unsafe_allow_html=True)
                        df_final = df_c[cols_slik3]
                        
                    elif sel_format == "slik 2 (Aldista)":
                        df_b = df_f.rename(columns={
                            "JENIS_MAPPED": "Jenis Penggunaan", 
                            "NAMA JASA KEUANGAN": "Bank/Lembaga pembiayaan",
                            "JENIS_ORIGINAL": "Jenis", 
                            "PLAFON_AWAL": "Plafon Awal",  
                            "BAKI DEBET": "OS (Rp)", 
                            "TGL_MULAI": "Tanggal Akad Akhir",
                            "JATUH_TEMPO": "Tanggal Jatuh Tempo",
                            "KOL_TERAKHIR": "Kol Terakhir",
                            "KOL_TERBURUK": "Kol terburuk",
                            "BUNGA": "Rate (%)",
                            "RESTRUK": "Restrukturisasi",
                            "KONDISI": "Kondisi"  
                        })
                        df_b["Jumlah Hari Kol"] = "-"
                        cols = ["NO", "Jenis Penggunaan", "Bank/Lembaga pembiayaan", "Jenis", "Plafon Awal", "OS (Rp)", "Tanggal Akad Akhir", "Tanggal Jatuh Tempo", "Kol Terakhir", "Kol terburuk", "Jumlah Hari Kol", "Rate (%)", "Restrukturisasi", "Kondisi"]
                        st.markdown('<div class="blue-header">', unsafe_allow_html=True); st.dataframe(df_b[cols], use_container_width=True, hide_index=True); st.markdown('</div>', unsafe_allow_html=True)
                        st.markdown(f"""<div style="background-color:#0000FF; color:white; padding:10px; font-weight:bold; text-align:center;">Total Outstanding: {format_rupiah(df_f['RAW_BAKI'].sum())}</div>""", unsafe_allow_html=True)
                        df_final = df_b[cols]
                        
                    else:
                        df_d = df_f.rename(columns={
                            "JENIS_ORIGINAL": "JENIS", 
                            "JENIS_MAPPED": "JENIS PENGGUNAAN",
                            "TGL_MULAI": "TGL AKAD AKHIR",
                            "JATUH_TEMPO": "TGL JATUH TEMPO",
                            "KOL_TERAKHIR": "KOL TERAKHIR", 
                            "KOL_TERBURUK": "KOL TERBURUK"
                        })
                        cols_slik1 = ["NO", "NAMA JASA KEUANGAN", "JENIS", "JENIS PENGGUNAAN", "PLAFON", "BAKI DEBET", "TGL AKAD AKHIR", "TGL JATUH TEMPO", "KOL TERAKHIR", "KOL TERBURUK", "BUNGA", "KONDISI"]
                        df_final = df_d[cols_slik1]
                        st.dataframe(df_final, use_container_width=True, hide_index=True)

                st.divider(); st.subheader("📥 Unduh Laporan")
                b1, b2, b3 = st.columns(3)
                id_i = {"nama": nama_v, "nik": nik_v, "alamat": alamat_v, "tgl": tgl_laporan_v, "tmpt_lahir": tmpt_lahir_v, "tgl_lahir": tgl_lahir_v, "jk": jk_v, "npwp": npwp_v, "pekerjaan": pekerjaan_v}
                aud_i = {"skor": skor_v, "plafon": format_rupiah(plafon_v), "baki": format_rupiah(baki_v), "util": f"{util_v:.2f}%", "total_kred": total_kred, "posisi": posisi_data_v}
                with b1: st.download_button("Excel (.xlsx)", icon="📊", data=export_excel(id_i, aud_i, df_final), file_name=f"Audit_{nama_v}.xlsx", key=f"xlsx_{uploaded_file.name}")
                with b2: st.download_button("Word (.docx)", icon="📝", data=export_word(id_i, aud_i, df_final), file_name=f"Audit_{nama_v}.docx", key=f"word_{uploaded_file.name}")
                with b3:
                    if st.button("Generate PDF", icon="⚙️", key=f"gen_{uploaded_file.name}"):
                        pdf_b = export_pdf(id_i, aud_i, df_final)
                        st.download_button("Klik Simpan PDF", icon="📕", data=pdf_b, file_name=f"Audit_{nama_v}.pdf", key=f"pdf_{uploaded_file.name}")

        except Exception as e: st.error(f"❌ Kesalahan pada file {uploaded_file.name}: {e}")
else: st.info("Unggah satu atau beberapa file .txt iDEB untuk memproses.")
